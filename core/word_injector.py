"""Inject placeholder values into a Word (.docx) template."""
from __future__ import annotations

import copy
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def inject(template_path: Path, output_path: Path, replacements: Dict[str, str | List[str]]) -> None:
	"""
	Load *template_path*, replace all {{MACRO}} placeholders, and save to *output_path*.

	replacements values can be:
	- str  → simple text replacement (handles fragmented runs)
	- List[str] → replace the paragraph that contains {{MACRO}} with one
				  paragraph per bullet, copying the original paragraph's XML format.
	"""
	doc = Document(str(template_path))
	output_path.parent.mkdir(parents=True, exist_ok=True)

	seen_cells: set = set()
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				if cell._tc not in seen_cells:
					seen_cells.add(cell._tc)
					_process_container(cell, replacements)

	_process_container(doc, replacements)

	doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _process_container(container, replacements: Dict[str, str | List[str]]) -> None:
	paragraphs = list(container.paragraphs)
	for para in paragraphs:
		# First, consolidate fragmented runs that together form a placeholder
		_consolidate_runs(para)

		full_text = para.text
		for macro, value in replacements.items():
			placeholder = f"{{{{{macro}}}}}"
			if placeholder not in full_text:
				continue

			if isinstance(value, list):
				_replace_paragraph_with_bullets(para, placeholder, value)
				break  # paragraph may be removed; stop iterating macros
			else:
				_replace_text_in_paragraph(para, placeholder, str(value))
				full_text = para.text  # update for next macro


def _consolidate_runs(para) -> None:
	"""
	Merge consecutive runs so that a placeholder like {{FULL_NAME}} is not split
	across multiple run objects (a common Word artifact).
	"""
	runs = para.runs
	if not runs:
		return

	combined = "".join(r.text for r in runs)
	# If the combined text contains a placeholder that no single run has,
	# consolidate all text into the first run and clear the rest.
	import re
	if re.search(r"\{\{[A-Z0-9_]+\}\}", combined):
		runs[0].text = combined
		for r in runs[1:]:
			r.text = ""


def _replace_text_in_paragraph(para, placeholder: str, value: str) -> None:
	for run in para.runs:
		if placeholder in run.text:
			run.text = run.text.replace(placeholder, value)


def _replace_paragraph_with_bullets(para, placeholder: str, bullets: List[str]) -> None:
	"""
	Replace the paragraph containing *placeholder* with N new paragraphs,
	one per bullet.

	Strategy: deep-copy the original paragraph element for each bullet
	(preserving its full XML — pPr, character styles, numPr, fonts, etc.),
	set the bullet text in the first run, and *remove* all extra runs from
	the copy entirely.

	Merely clearing extra runs' text (r.text = "") leaves empty <w:r> nodes
	in the XML. When multiple structurally-identical paragraphs are inserted
	consecutively, Word uses those ghost runs to resolve character formatting
	and ends up dropping the bullet paragraph style on every paragraph after
	the first. Removing them gives each bullet a clean, independent structure
	that Word renders consistently.
	"""
	parent = para._element.getparent()
	if parent is None:
		return

	idx = list(parent).index(para._element)

	new_elements = []
	for bullet_text in bullets:
		# Full deep-copy preserves every XML attribute, namespace, pPr, rPr,
		# character style, numPr — the complete formatting of the original.
		new_p = copy.deepcopy(para._element)

		runs_xml = new_p.findall(f".//{qn('w:r')}")
		if runs_xml:
			# Write the bullet text into the first run's <w:t>
			t_elem = runs_xml[0].find(qn("w:t"))
			if t_elem is None:
				t_elem = etree.SubElement(runs_xml[0], qn("w:t"))
			t_elem.text = bullet_text
			t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

			# Remove every extra run completely — do NOT just clear their text.
			# Ghost empty runs cause Word to mis-render bullet formatting on
			# the 2nd+ paragraphs when consecutive deep-copies are present.
			for r in runs_xml[1:]:
				r_parent = r.getparent()
				if r_parent is not None:
					r_parent.remove(r)
		else:
			# Paragraph had no runs at all — create a minimal one.
			new_r = etree.SubElement(new_p, qn("w:r"))
			t_elem = etree.SubElement(new_r, qn("w:t"))
			t_elem.text = bullet_text
			t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

		new_elements.append(new_p)

	# Insert all bullet paragraphs, then remove the original placeholder line.
	for offset, elem in enumerate(new_elements):
		parent.insert(idx + offset, elem)
	parent.remove(para._element)
